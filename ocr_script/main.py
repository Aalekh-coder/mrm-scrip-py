import os
import re
import sys
import json



PDF_MAX_PAGES  = 5    # max matched pages to return per PDF file
EXCEL_MAX_ROWS = 10   # max matched rows to return per Excel file
EXCEL_MAX_COLS = 10   # max columns to show per Excel row (rest are truncated)


MISSING = []
try:
    import pdfplumber
except ImportError:
    MISSING.append("pdfplumber")

try:
    import openpyxl
except ImportError:
    MISSING.append("openpyxl")

try:
    from docx import Document as DocxDocument
except ImportError:
    MISSING.append("python-docx")

if MISSING:
    print("Missing dependencies. Install with:")
    print(f"  pip install {' '.join(MISSING)}")
    sys.exit(1)


# ---------------------------------------------------------------------------
# Extractors — each returns a list of { "source": str, "text": str }
# ---------------------------------------------------------------------------

def extract_from_pdf(filepath: str) -> list[dict]:
    """
    Extract one chunk per page (full page text merged into a single string).
    Table rows are also extracted as separate chunks.
    Deduplication happens at the page level — one result per matched page.
    """
    chunks = []
    try:
        with pdfplumber.open(filepath) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):

                # Full page text → single chunk (one summary per page)
                text = (page.extract_text() or "").strip()
                if text:
                    chunks.append({
                        "source": f"{os.path.basename(filepath)} — Page {page_num}",
                        "text":   text,
                        "_page":  page_num,   # used for dedup + limit
                    })

                # Tables — one chunk per row
                tables = page.extract_tables() or []
                for t_idx, table in enumerate(tables, start=1):
                    for r_idx, row in enumerate(table, start=1):
                        row_text = " | ".join(cell or "" for cell in row).strip()
                        if row_text:
                            chunks.append({
                                "source": (
                                    f"{os.path.basename(filepath)} — "
                                    f"Page {page_num}, Table {t_idx}, Row {r_idx}"
                                ),
                                "text":  row_text,
                                "_page": page_num,
                            })
    except Exception as e:
        print(f"  [PDF Error] {filepath}: {e}")
    return chunks


def extract_from_excel(filepath: str) -> list[dict]:
    """
    Extract cell content from every sheet.
    Columns beyond EXCEL_MAX_COLS are replaced with a '… (+N more)' note.
    """
    chunks = []
    try:
        wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            for row_idx, row in enumerate(ws.iter_rows(values_only=True), start=1):
                all_cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if not all_cells:
                    continue

                # ✏️ EXCEL_MAX_COLS: cap visible columns, note how many are hidden
                if len(all_cells) > EXCEL_MAX_COLS:
                    hidden   = len(all_cells) - EXCEL_MAX_COLS
                    all_cells = all_cells[:EXCEL_MAX_COLS] + [f"… (+{hidden} more cols)"]

                row_text = " | ".join(all_cells)
                chunks.append({
                    "source": f"{os.path.basename(filepath)} — Sheet '{sheet_name}', Row {row_idx}",
                    "text":   row_text,
                })
        wb.close()
    except Exception as e:
        print(f"  [Excel Error] {filepath}: {e}")
    return chunks


def extract_from_docx(filepath: str) -> list[dict]:
    """
    Extract paragraphs and table cells from a Word document.
    """
    chunks = []
    try:
        doc = DocxDocument(filepath)

        # Paragraphs
        for para_idx, para in enumerate(doc.paragraphs, start=1):
            text = para.text.strip()
            if text:
                chunks.append({
                    "source": f"{os.path.basename(filepath)} — Paragraph {para_idx}",
                    "text":   text,
                })

        # Tables
        for t_idx, table in enumerate(doc.tables, start=1):
            for r_idx, row in enumerate(table.rows, start=1):
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    chunks.append({
                        "source": f"{os.path.basename(filepath)} — Table {t_idx}, Row {r_idx}",
                        "text":   row_text,
                    })
    except Exception as e:
        print(f"  [DOCX Error] {filepath}: {e}")
    return chunks


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------

EXTRACTOR_MAP = {
    ".pdf":  extract_from_pdf,
    ".xlsx": extract_from_excel,
    ".xls":  extract_from_excel,
    ".docx": extract_from_docx,
    ".doc":  extract_from_docx,
}


def extract_chunks(filepath: str) -> list[dict]:
    ext = os.path.splitext(filepath)[1].lower()
    extractor = EXTRACTOR_MAP.get(ext)
    if not extractor:
        print(f"  [Skip] Unsupported file type: {filepath}")
        return []
    return extractor(filepath)


# ---------------------------------------------------------------------------
# Search
# ---------------------------------------------------------------------------

def search_chunks(chunks: list[dict], keyword: str, case_sensitive: bool = False) -> list[dict]:
    """Return all chunks whose text contains keyword (exact match)."""
    flag    = 0 if case_sensitive else re.IGNORECASE
    pattern = re.compile(re.escape(keyword), flag)

    matches = []
    for chunk in chunks:
        if pattern.search(chunk["text"]):
            matches.append(chunk)
    return matches


# ---------------------------------------------------------------------------
# Summariser
# ---------------------------------------------------------------------------

def _summarise(text: str, keyword: str, context_words: int = 12) -> str:
    """
    Return a crisp snippet with ~context_words on each side of the keyword.
    Text <= 120 chars is returned as-is.
    """
    text = text.strip()
    if len(text) <= 120:
        return text

    pattern = re.compile(re.escape(keyword), re.IGNORECASE)
    m = pattern.search(text)
    if not m:
        return text[:120] + "…"

    words  = text.split()
    pos    = 0
    kw_idx = 0
    for idx, word in enumerate(words):
        pos += len(word) + 1
        if pos > m.start():
            kw_idx = idx
            break

    start   = max(0, kw_idx - context_words)
    end     = min(len(words), kw_idx + context_words + 1)
    snippet = " ".join(words[start:end])
    if start > 0:
        snippet = "…" + snippet
    if end < len(words):
        snippet = snippet + "…"
    return snippet


# ---------------------------------------------------------------------------
# Per-file limiters
# ---------------------------------------------------------------------------

def _apply_pdf_limit(matches: list[dict], filepath: str) -> list[dict]:
    """
    Keep only the first PDF_MAX_PAGES distinct matched pages.
    Within a kept page, only the first (full-page) chunk is returned —
    table row chunks on the same page are skipped to avoid repetition.
    """
    ext = os.path.splitext(filepath)[1].lower()
    if ext != ".pdf":
        return matches

    seen_pages  = []   # ordered list of page numbers already included
    filtered    = []

    for chunk in matches:
        page = chunk.get("_page")
        if page is None:
            filtered.append(chunk)
            continue

        if page not in seen_pages:
            if len(seen_pages) >= PDF_MAX_PAGES:
                continue                      # ✏️ PDF_MAX_PAGES limit reached
            seen_pages.append(page)
            filtered.append(chunk)
        # else: same page already represented — skip duplicate

    return filtered


def _apply_excel_limit(matches: list[dict], filepath: str) -> list[dict]:
    """
    Keep only the first EXCEL_MAX_ROWS matched rows per Excel file.
    """
    ext = os.path.splitext(filepath)[1].lower()
    if ext not in (".xlsx", ".xls"):
        return matches
    return matches[:EXCEL_MAX_ROWS]   # ✏️ EXCEL_MAX_ROWS limit


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def search_files(
    keyword: str,
    filepaths: list[str],
    case_sensitive: bool = False,
) -> dict:
    """
    Search keyword across all given files.

    Returns:
    {
        "keyword": "TypeScript",
        "total_matches": 5,
        "results": [
            {
                "file_name": "deep-dive.pdf",
                "file_path": "C:\\docs\\deep-dive.pdf",
                "location":  "Page 330",
                "summary":   "…TypeScript Compiler Internals File: System system.ts…"
            },
            ...
        ]
    }
    """
    all_results = []

    for fp in filepaths:
        if not os.path.isfile(fp):
            print(f"  [Not Found] {fp}")
            continue

        print(f"  Reading: {fp}")
        ext     = os.path.splitext(fp)[1].lower()
        chunks  = extract_chunks(fp)
        matches = search_chunks(chunks, keyword, case_sensitive)

        # Apply per-type limits
        matches = _apply_pdf_limit(matches, fp)
        matches = _apply_excel_limit(matches, fp)

        for m in matches:
            source_parts = m["source"].split(" — ", 1)
            location     = source_parts[1] if len(source_parts) > 1 else m["source"]
            summary      = _summarise(m["text"], keyword)

            all_results.append({
                "file_name": os.path.basename(fp),
                "file_path": os.path.abspath(fp),
                "location":  location,
                "summary":   summary,
            })

    return {
        "keyword":       keyword,
        "total_matches": len(all_results),
        "results":       all_results,
    }


# ---------------------------------------------------------------------------
# Pretty printer
# ---------------------------------------------------------------------------

def print_results(output: dict):
    print("\n" + json.dumps(output, indent=2, ensure_ascii=False) + "\n")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def _collect_files_from_user() -> list[str]:
    supported_exts = set(EXTRACTOR_MAP.keys())
    files = []

    print("\nEnter file or folder paths (one per line). Press ENTER on empty line when done.")
    while True:
        path = input("  Path: ").strip().strip('"').strip("'")
        if not path:
            break
        if os.path.isdir(path):
            found = [
                os.path.join(path, f)
                for f in os.listdir(path)
                if os.path.splitext(f)[1].lower() in supported_exts
            ]
            if found:
                print(f"    Found {len(found)} supported file(s) in folder.")
                files.extend(found)
            else:
                print("    No supported files found in that folder.")
        elif os.path.isfile(path):
            files.append(path)
        else:
            print("    Path not found — skipping.")

    return files


if __name__ == "__main__":
    print("\n╔══════════════════════════════════╗")
    print("║       OCR Keyword Search         ║")
    print("║  PDF · Excel · Word              ║")
    print("╚══════════════════════════════════╝")

    keyword = input("\nEnter keyword to search: ").strip()
    if not keyword:
        print("No keyword entered. Exiting.")
        sys.exit(0)

    case_input     = input("Case-sensitive? (y/N): ").strip().lower()
    case_sensitive = case_input == "y"

    filepaths = _collect_files_from_user()
    if not filepaths:
        print("No files to search. Exiting.")
        sys.exit(0)

    print(f"\nSearching {len(filepaths)} file(s) for \"{keyword}\"...\n")
    output = search_files(keyword, filepaths, case_sensitive=case_sensitive)
    print_results(output)